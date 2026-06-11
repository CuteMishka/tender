from __future__ import annotations

import io
import unittest

from docx import Document
from openpyxl import Workbook

from app.document_extract import extract_text_from_bytes


class DocumentExtractTests(unittest.TestCase):
    def test_plain_text_and_csv(self) -> None:
        self.assertEqual(extract_text_from_bytes("spec.txt", "Срок 30 дней".encode()), "Срок 30 дней")
        self.assertIn("Цена | 1000", extract_text_from_bytes("price.csv", "Цена,1000".encode()))

    def test_docx_paragraphs_and_tables(self) -> None:
        document = Document()
        document.add_paragraph("Техническая спецификация")
        table = document.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "Срок"
        table.cell(0, 1).text = "45 дней"
        buffer = io.BytesIO()
        document.save(buffer)
        text = extract_text_from_bytes("spec.docx", buffer.getvalue())
        self.assertIn("Техническая спецификация", text)
        self.assertIn("Срок | 45 дней", text)

    def test_xlsx_sheets_and_cells(self) -> None:
        workbook = Workbook()
        sheet = workbook.active
        sheet.title = "Условия"
        sheet.append(["Бюджет", 2500000])
        buffer = io.BytesIO()
        workbook.save(buffer)
        text = extract_text_from_bytes("terms.xlsx", buffer.getvalue())
        self.assertIn("Лист: Условия", text)
        self.assertIn("Бюджет | 2500000", text)

    def test_unsupported_format_has_clear_error(self) -> None:
        with self.assertRaisesRegex(ValueError, "не поддерживается"):
            extract_text_from_bytes("archive.zip", b"not-a-zip")


if __name__ == "__main__":
    unittest.main()
