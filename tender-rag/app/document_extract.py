"""Extract text from document formats used in tender lots."""

from __future__ import annotations

import csv
import hashlib
import io
import json
import os
import re
import shutil
import subprocess
import tempfile
import zipfile
from collections import OrderedDict
from html.parser import HTMLParser
from pathlib import Path
from threading import Lock
from xml.etree import ElementTree as ET

from pypdf import PdfReader

XML_NS = "{http://schemas.openxmlformats.org/spreadsheetml/2006/main}"
OCR_ENABLED = os.environ.get("DOCUMENT_OCR_ENABLED", "true").strip().lower() not in {"0", "false", "no", "off"}
OCR_MAX_PAGES = int(os.environ.get("DOCUMENT_OCR_MAX_PAGES", "25"))
OCR_DPI = int(os.environ.get("DOCUMENT_OCR_DPI", "200"))
OCR_LANGS = tuple(
    lang.strip()
    for lang in os.environ.get("DOCUMENT_OCR_LANGS", "rus+kaz+eng,rus+eng,eng").split(",")
    if lang.strip()
)
MIN_MEANINGFUL_TEXT_CHARS = 80

# Извлечение текста из PDF/DOCX — дорогая операция. Cloudy переотправляет
# одни и те же документы на каждый вопрос диалога, поэтому кэшируем результат
# по sha256 содержимого: одинаковые байты всегда дают одинаковый текст.
_EXTRACT_CACHE_MAX = 256
_extract_cache: "OrderedDict[str, str]" = OrderedDict()
_extract_cache_lock = Lock()


def _extract_cache_get(key: str) -> str | None:
    with _extract_cache_lock:
        value = _extract_cache.get(key)
        if value is not None:
            _extract_cache.move_to_end(key)
        return value


def _extract_cache_set(key: str, value: str) -> None:
    with _extract_cache_lock:
        _extract_cache[key] = value
        _extract_cache.move_to_end(key)
        while len(_extract_cache) > _EXTRACT_CACHE_MAX:
            _extract_cache.popitem(last=False)


def _norm(s: str) -> str:
    t = s.replace("\r\n", "\n").replace("\r", "\n")
    t = re.sub(r"\n{3,}", "\n\n", t)
    return t.strip()


def _decode_text(data: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "cp1251", "koi8-r", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="ignore")


def _looks_like_text(text: str) -> bool:
    sample = text[:8000]
    if not sample.strip():
        return False
    control = sum(
        1
        for ch in sample
        if ord(ch) < 32 and ch not in "\n\r\t\f\b"
    )
    return control / max(1, len(sample)) < 0.05


def _has_meaningful_text(text: str) -> bool:
    return sum(1 for ch in text if ch.isalnum()) >= MIN_MEANINGFUL_TEXT_CHARS


def _extract_text_from_pdf_pypdf(data: bytes) -> str:
    reader = PdfReader(io.BytesIO(data))
    parts: list[str] = []
    for page in reader.pages:
        parts.append(page.extract_text() or "")
    return _norm("\n".join(parts))


def _extract_text_from_pdf_pymupdf(data: bytes) -> str:
    try:
        import fitz  # PyMuPDF
    except Exception:
        return ""
    try:
        with fitz.open(stream=data, filetype="pdf") as doc:
            return _norm("\n".join(page.get_text("text") or "" for page in doc))
    except Exception:
        return ""


def _extract_text_from_pdf_ocr(data: bytes) -> str:
    if not OCR_ENABLED or OCR_MAX_PAGES <= 0:
        return ""
    try:
        from pdf2image import convert_from_bytes
        import pytesseract
    except Exception:
        return ""

    try:
        pages = convert_from_bytes(
            data,
            dpi=OCR_DPI,
            first_page=1,
            last_page=OCR_MAX_PAGES,
            fmt="png",
            thread_count=1,
        )
    except Exception:
        return ""

    for langs in OCR_LANGS or ("eng",):
        parts: list[str] = []
        try:
            for page in pages:
                parts.append(pytesseract.image_to_string(page, lang=langs) or "")
        except Exception:
            continue
        text = _norm("\n".join(parts))
        if text:
            return text
    return ""


def _extract_text_from_pdf(data: bytes) -> str:
    texts: list[str] = []
    for extractor in (_extract_text_from_pdf_pypdf, _extract_text_from_pdf_pymupdf):
        try:
            text = extractor(data)
        except Exception:
            text = ""
        if text:
            texts.append(text)

    best = max(texts, key=len, default="")
    if _has_meaningful_text(best):
        return best

    ocr_text = _extract_text_from_pdf_ocr(data)
    if _has_meaningful_text(ocr_text):
        return ocr_text
    return ocr_text or best


def _extract_text_from_docx(data: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(data))
    parts: list[str] = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return _norm("\n".join(parts))


def _extract_text_from_pptx(data: bytes) -> str:
    parts: list[str] = []
    with zipfile.ZipFile(io.BytesIO(data)) as zf:
        slide_names = sorted(
            name
            for name in zf.namelist()
            if name.startswith("ppt/slides/slide") and name.endswith(".xml")
        )
        for slide_name in slide_names:
            root = ET.fromstring(zf.read(slide_name))
            for node in root.iter():
                if node.tag.endswith("}t") and node.text and node.text.strip():
                    parts.append(node.text.strip())
    return _norm("\n".join(parts))


def _extract_text_from_rtf(data: bytes) -> str:
    raw = _decode_text(data)
    try:
        from striprtf.striprtf import rtf_to_text

        return _norm(rtf_to_text(raw))
    except Exception:
        return _extract_text_from_text(data)


def _extract_text_from_xls(data: bytes) -> str:
    try:
        import xlrd

        book = xlrd.open_workbook(file_contents=data)
        rows: list[str] = []
        for sheet in book.sheets():
            for row_index in range(sheet.nrows):
                cells = [
                    str(sheet.cell_value(row_index, col_index)).strip()
                    for col_index in range(sheet.ncols)
                ]
                cells = [cell for cell in cells if cell]
                if cells:
                    rows.append(" | ".join(cells))
        return _norm("\n".join(rows))
    except Exception:
        return ""


def _extract_text_via_command(command: str, filename: str, data: bytes) -> str | None:
    exe = shutil.which(command)
    if not exe:
        return None
    suffix = Path(filename or "document").suffix or ".bin"
    with tempfile.TemporaryDirectory() as tmpdir:
        src = Path(tmpdir) / f"source{suffix}"
        src.write_bytes(data)
        try:
            proc = subprocess.run(
                [exe, str(src)],
                capture_output=True,
                timeout=45,
                check=False,
            )
        except (OSError, subprocess.TimeoutExpired):
            return None
    if proc.returncode != 0:
        return None
    return _norm(_decode_text(proc.stdout))


def _extract_text_from_text(data: bytes) -> str:
    text = _decode_text(data)
    if not _looks_like_text(text):
        return ""
    return _norm(text)


def _extract_text_from_delimited(data: bytes, delimiter: str) -> str:
    rows: list[str] = []
    reader = csv.reader(io.StringIO(_decode_text(data)), delimiter=delimiter)
    for row in reader:
        cells = [cell.strip() for cell in row if cell and cell.strip()]
        if cells:
            rows.append(" | ".join(cells))
    return _norm("\n".join(rows))


class _HTMLTextExtractor(HTMLParser):
    block_tags = {
        "p",
        "div",
        "br",
        "li",
        "tr",
        "table",
        "section",
        "article",
        "header",
        "footer",
        "h1",
        "h2",
        "h3",
        "h4",
        "h5",
        "h6",
    }

    def __init__(self) -> None:
        super().__init__()
        self.parts: list[str] = []

    def handle_starttag(self, tag: str, _attrs):
        if tag.lower() in self.block_tags:
            self.parts.append("\n")

    def handle_endtag(self, tag: str):
        if tag.lower() in self.block_tags:
            self.parts.append("\n")

    def handle_data(self, data: str):
        text = data.strip()
        if text:
            self.parts.append(text)


def _extract_text_from_html(data: bytes) -> str:
    parser = _HTMLTextExtractor()
    parser.feed(_decode_text(data))
    return _norm(" ".join(parser.parts))


def _extract_text_from_json(data: bytes) -> str:
    text = _decode_text(data).strip()
    if not text:
        return ""
    parsed = json.loads(text)
    return _norm(json.dumps(parsed, ensure_ascii=False, indent=2))


def _extract_text_from_xml(data: bytes) -> str:
    root = ET.fromstring(_decode_text(data).encode("utf-8"))
    parts: list[str] = []

    def visit(node: ET.Element) -> None:
        text = " ".join((node.text or "").split())
        if text:
            parts.append(text)
        for child in list(node):
            visit(child)
            tail = " ".join((child.tail or "").split())
            if tail:
                parts.append(tail)

    visit(root)
    return _norm("\n".join(parts))


def _xlsx_shared_strings(zf: zipfile.ZipFile) -> list[str]:
    try:
        raw = zf.read("xl/sharedStrings.xml")
    except KeyError:
        return []
    root = ET.fromstring(raw)
    values: list[str] = []
    for si in root.iter(f"{XML_NS}si"):
        text_parts = [node.text or "" for node in si.iter(f"{XML_NS}t")]
        values.append(_norm("".join(text_parts)))
    return values


def _extract_text_from_xlsx(data: bytes) -> str:
    with zipfile.ZipFile(io.BytesIO(data)) as zf:
        shared = _xlsx_shared_strings(zf)
        parts: list[str] = []
        sheet_names = sorted(
            name
            for name in zf.namelist()
            if name.startswith("xl/worksheets/sheet") and name.endswith(".xml")
        )
        for sheet_name in sheet_names:
            root = ET.fromstring(zf.read(sheet_name))
            for row in root.iter(f"{XML_NS}row"):
                cells: list[str] = []
                for cell in row.findall(f"{XML_NS}c"):
                    cell_type = cell.attrib.get("t")
                    text = ""
                    value = cell.find(f"{XML_NS}v")
                    if cell_type == "s" and value is not None and value.text is not None:
                        try:
                            idx = int(value.text)
                        except ValueError:
                            idx = -1
                        if 0 <= idx < len(shared):
                            text = shared[idx]
                    elif cell_type == "inlineStr":
                        inline = cell.find(f"{XML_NS}is")
                        if inline is not None:
                            text = "".join(part.text or "" for part in inline.iter(f"{XML_NS}t"))
                    elif value is not None and value.text is not None:
                        text = value.text
                    text = text.strip()
                    if text:
                        cells.append(text)
                if cells:
                    parts.append(" | ".join(cells))
        return _norm("\n".join(parts))


def _zip_office_format(data: bytes) -> str:
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as zf:
            names = set(zf.namelist())
    except zipfile.BadZipFile:
        return ""
    if "word/document.xml" in names:
        return ".docx"
    if "xl/workbook.xml" in names or any(
        name.startswith("xl/worksheets/") and name.endswith(".xml") for name in names
    ):
        return ".xlsx"
    if "ppt/presentation.xml" in names:
        return ".pptx"
    return ".zip"


def _detect_document_format(filename: str, data: bytes) -> str:
    head = data[:4096].lstrip()
    if data.startswith(b"%PDF"):
        return ".pdf"
    if head.startswith(b"{\\rtf"):
        return ".rtf"
    if zipfile.is_zipfile(io.BytesIO(data)):
        return _zip_office_format(data)

    text_head = _decode_text(head).lstrip().lower()
    if text_head.startswith("<!doctype html") or text_head.startswith("<html") or "<html" in text_head[:300]:
        return ".html"
    if text_head.startswith("{") or text_head.startswith("["):
        return ".json"
    if text_head.startswith("<?xml") or text_head.startswith("<"):
        return ".xml"

    ext = Path(filename or "").suffix.lower()
    return ext


def _extract_text_via_office(filename: str, data: bytes) -> str | None:
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        return None

    suffix = Path(filename or "document").suffix or ".bin"
    with tempfile.TemporaryDirectory() as tmpdir:
        src = Path(tmpdir) / f"source{suffix}"
        src.write_bytes(data)
        outdir = Path(tmpdir) / "out"
        outdir.mkdir(parents=True, exist_ok=True)
        try:
            proc = subprocess.run(
                [
                    soffice,
                    "--headless",
                    "--convert-to",
                    "txt:Text",
                    "--outdir",
                    str(outdir),
                    str(src),
                ],
                capture_output=True,
                timeout=45,
                check=False,
            )
        except (OSError, subprocess.TimeoutExpired):
            return None
        if proc.returncode != 0:
            return None
        converted = outdir / f"{src.stem}.txt"
        if not converted.exists():
            txt_files = sorted(outdir.glob("*.txt"))
            if not txt_files:
                return None
            converted = txt_files[0]
        try:
            return _norm(converted.read_text(encoding="utf-8", errors="ignore"))
        except OSError:
            return None


def _extract_text_from_plain_or_unknown(filename: str, data: bytes) -> str:
    text = _extract_text_from_text(data)
    if text:
        return text
    fallback = _extract_text_via_office(filename, data)
    if fallback:
        return fallback
    return ""


def _extract_text_from_doc(filename: str, data: bytes) -> str:
    for command in ("antiword", "catdoc"):
        text = _extract_text_via_command(command, filename, data)
        if text:
            return text
    fallback = _extract_text_via_office(filename, data)
    if fallback:
        return fallback
    return _extract_text_from_plain_or_unknown(filename, data)


def _extract_text_from_legacy_office(filename: str, data: bytes) -> str:
    fallback = _extract_text_via_office(filename, data)
    if fallback:
        return fallback
    return _extract_text_from_plain_or_unknown(filename, data)


def extract_text_from_bytes(filename: str, data: bytes) -> str:
    if not data:
        return ""
    doc_format = _detect_document_format(filename, data)
    cache_key = f"{hashlib.sha256(data).hexdigest()}:{doc_format}"
    cached = _extract_cache_get(cache_key)
    if cached is not None:
        return cached
    result = _extract_text_from_bytes_uncached(filename, data, doc_format)
    _extract_cache_set(cache_key, result)
    return result


def _extract_text_from_bytes_uncached(filename: str, data: bytes, doc_format: str) -> str:
    if doc_format == ".pdf":
        return _extract_text_from_pdf(data)
    if doc_format == ".docx":
        try:
            return _extract_text_from_docx(data)
        except Exception:
            return _extract_text_from_legacy_office(filename, data)
    if doc_format in {".txt", ".md", ".log"}:
        return _extract_text_from_text(data)
    if doc_format == ".csv":
        return _extract_text_from_delimited(data, ",")
    if doc_format == ".tsv":
        return _extract_text_from_delimited(data, "\t")
    if doc_format in {".html", ".htm"}:
        return _extract_text_from_html(data)
    if doc_format == ".json":
        try:
            return _extract_text_from_json(data)
        except Exception:
            return _extract_text_from_text(data)
    if doc_format == ".xml":
        try:
            return _extract_text_from_xml(data)
        except Exception:
            return _extract_text_from_text(data)
    if doc_format == ".xlsx":
        try:
            return _extract_text_from_xlsx(data)
        except Exception:
            fallback = _extract_text_via_office(filename, data)
            if fallback:
                return fallback
            return ""
    if doc_format == ".xls":
        text = _extract_text_from_xls(data)
        if text:
            return text
        return _extract_text_from_legacy_office(filename, data)
    if doc_format == ".rtf":
        text = _extract_text_from_rtf(data)
        if text:
            return text
        return _extract_text_from_legacy_office(filename, data)
    if doc_format == ".doc":
        return _extract_text_from_doc(filename, data)
    if doc_format == ".pptx":
        try:
            text = _extract_text_from_pptx(data)
        except Exception:
            text = ""
        if text:
            return text
        return _extract_text_from_legacy_office(filename, data)
    if doc_format in {".odt", ".ods", ".ppt"}:
        return _extract_text_from_legacy_office(filename, data)
    return _extract_text_from_plain_or_unknown(filename, data)
