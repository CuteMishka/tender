from __future__ import annotations

import os
from datetime import datetime
from decimal import Decimal, ROUND_HALF_UP
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document
from docxtpl import DocxTemplate
from sqlalchemy import func, select
from sqlalchemy.ext.asyncio import AsyncSession

from app.models import ClientProfile, CommercialProposal
from app.services.audit_service import log_lot_event
from app.services.lot_service import get_or_import_lot
from app.services.storage_service import ObjectStorageService

PACKAGE_BASE_PRICES = {
    "phishing": Decimal("350000"),
    "awareness": Decimal("500000"),
    "courses": Decimal("750000"),
    "citizensec_standard": Decimal("1200000"),
    "citizensec_enterprise": Decimal("2500000"),
}
VAT_RATE = Decimal(os.environ.get("CP_VAT_RATE", "0.12"))
MARGIN_RATE = Decimal(os.environ.get("CP_MARGIN_RATE", "0.25"))


async def generate_commercial_proposal(
    session: AsyncSession,
    lot_id: str,
    service_package: str,
    discount_percent: float = 0,
    client_payload: dict[str, Any] | None = None,
    parameters: dict[str, Any] | None = None,
    created_by_user_id: int | None = None,
) -> CommercialProposal:
    lot = await get_or_import_lot(session, lot_id)
    client = await _upsert_client_profile(session, client_payload or {})
    price_payload = calculate_price(service_package, discount_percent, parameters or {}, client)
    proposal_number = await _get_or_create_proposal_number(session, lot.id)
    version = await _next_version(session, proposal_number)
    storage_key = f"commercial-proposals/{proposal_number}/v{version}.docx"
    docx_bytes = render_docx(
        template_name=str((parameters or {}).get("template_name") or "commercial_proposal.docx"),
        context={
            "proposal_number": proposal_number,
            "version": version,
            "date": datetime.utcnow().strftime("%d.%m.%Y"),
            "lot": lot,
            "client": client,
            "service_package": service_package,
            "discount_percent": discount_percent,
            "price": price_payload,
            "parameters": parameters or {},
        },
    )
    storage = await ObjectStorageService().put_bytes(
        storage_key,
        docx_bytes,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    proposal = CommercialProposal(
        lot_id=lot.id,
        client_profile_id=client.id if client else None,
        proposal_number=proposal_number,
        version=version,
        service_package=service_package,
        discount_percent=discount_percent,
        price_payload=price_payload,
        vat_included=True,
        total_amount=Decimal(str(price_payload["total_with_vat"])),
        storage_backend=str(storage["backend"]),
        storage_bucket=storage.get("bucket"),
        storage_key=str(storage["key"]),
        file_url=storage.get("url"),
        template_name=str((parameters or {}).get("template_name") or "commercial_proposal.docx"),
        created_by_user_id=created_by_user_id,
    )
    session.add(proposal)
    await session.flush()
    await log_lot_event(
        session,
        lot.id,
        "commercial_proposal_generated",
        f"Сформировано КП {proposal_number} v{version}",
        {"proposal_id": proposal.id, "storage_key": proposal.storage_key, "total_amount": str(proposal.total_amount)},
        created_by_user_id,
    )
    await session.commit()
    return proposal


def calculate_price(
    service_package: str,
    discount_percent: float,
    parameters: dict[str, Any],
    client: ClientProfile | None,
) -> dict[str, Any]:
    base = PACKAGE_BASE_PRICES.get(service_package, PACKAGE_BASE_PRICES["citizensec_standard"])
    employees = Decimal(str(parameters.get("employees_count") or (client.employees_count if client else 0) or 0))
    employee_component = employees * Decimal("1200")
    subtotal = base + employee_component
    margin = (subtotal * MARGIN_RATE).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
    before_discount = subtotal + margin
    discount = (before_discount * Decimal(str(discount_percent)) / Decimal("100")).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
    after_discount = before_discount - discount
    vat = (after_discount * VAT_RATE).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
    total = (after_discount + vat).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
    return {
        "currency": "KZT",
        "base": str(base),
        "employees_count": int(employees),
        "employee_component": str(employee_component),
        "margin_rate": str(MARGIN_RATE),
        "margin": str(margin),
        "discount_percent": discount_percent,
        "discount": str(discount),
        "vat_rate": str(VAT_RATE),
        "vat": str(vat),
        "total_with_vat": str(total),
    }


def render_docx(template_name: str, context: dict[str, Any]) -> bytes:
    template_dir = Path(os.environ.get("CP_TEMPLATE_DIR", "files/templates")).resolve()
    template_path = template_dir / template_name
    serialized = _serialize_docx_context(context)
    if template_path.is_file():
        doc = DocxTemplate(str(template_path))
        doc.render(serialized)
        buffer = BytesIO()
        doc.save(buffer)
        return buffer.getvalue()
    return _render_default_docx(serialized)


def _render_default_docx(context: dict[str, Any]) -> bytes:
    doc = Document()
    doc.add_heading("???????????? ???????????", level=1)
    doc.add_paragraph(f"?????: {context['proposal_number']} v{context['version']}")
    doc.add_paragraph(f"????: {context['date']}")
    doc.add_paragraph(f"???: {context['lot_id']} ? {context['lot_title']}")
    doc.add_paragraph(f"??????: {context['client_company_name']} ??? {context['client_bin']}")
    doc.add_paragraph(f"????? ?????: {context['service_package']}")
    doc.add_paragraph(f"??????: {context['discount_percent']}%")
    doc.add_paragraph(f"????? ? ???: {context['price']['total_with_vat']} KZT")
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


async def _upsert_client_profile(session: AsyncSession, payload: dict[str, Any]) -> ClientProfile | None:
    if not payload:
        return None
    bin_value = str(payload.get("bin") or payload.get("iin_bin") or "").strip() or None
    client = None
    if bin_value:
        result = await session.execute(select(ClientProfile).where(ClientProfile.bin == bin_value))
        client = result.scalar_one_or_none()
    if client is None:
        client = ClientProfile(bin=bin_value)
        session.add(client)
    client.company_name = payload.get("company_name") or payload.get("name") or client.company_name
    client.domain = payload.get("domain") or client.domain
    employees = payload.get("employees_count")
    if employees is not None:
        client.employees_count = int(employees)
    client.contacts = payload.get("contacts") or client.contacts
    await session.flush()
    return client


async def _get_or_create_proposal_number(session: AsyncSession, lot_id: str) -> str:
    result = await session.execute(
        select(CommercialProposal.proposal_number)
        .where(CommercialProposal.lot_id == lot_id)
        .order_by(CommercialProposal.created_at.asc())
        .limit(1)
    )
    existing = result.scalar_one_or_none()
    if existing:
        return existing
    sequence = await session.execute(select(func.count(CommercialProposal.id)))
    next_number = int(sequence.scalar_one() or 0) + 1
    return f"CP-{datetime.utcnow():%Y}-{next_number:05d}"


async def _next_version(session: AsyncSession, proposal_number: str) -> int:
    result = await session.execute(select(func.coalesce(func.max(CommercialProposal.version), 0)).where(CommercialProposal.proposal_number == proposal_number))
    return int(result.scalar_one() or 0) + 1


def _serialize_docx_context(context: dict[str, Any]) -> dict[str, Any]:
    lot = context["lot"]
    client = context.get("client")
    return {
        "proposal_number": context["proposal_number"],
        "version": context["version"],
        "date": context["date"],
        "lot_id": lot.id,
        "lot_title": lot.title or "",
        "lot_amount": str(lot.amount or ""),
        "lot_deadline": lot.deadline.strftime("%d.%m.%Y") if lot.deadline else "",
        "client_bin": client.bin if client else "",
        "client_company_name": client.company_name if client else "",
        "client_domain": client.domain if client else "",
        "employees_count": client.employees_count if client and client.employees_count else context["price"].get("employees_count", 0),
        "service_package": context["service_package"],
        "discount_percent": context["discount_percent"],
        "price": context["price"],
        "parameters": context["parameters"],
    }
