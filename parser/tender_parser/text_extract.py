from io import BytesIO
from pathlib import Path

from docx import Document
from pypdf import PdfReader


def extract_text_from_bytes(filename: str, data: bytes) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf(data)
    if suffix == ".docx":
        return _extract_docx(data)
    if suffix in {".txt", ".csv", ".html", ".htm"}:
        return data.decode("utf-8", errors="ignore")
    if suffix == ".doc":
        return data.decode("utf-8", errors="ignore")
    return data.decode("utf-8", errors="ignore")


def _extract_pdf(data: bytes) -> str:
    reader = PdfReader(BytesIO(data))
    pages = [(page.extract_text() or "").strip() for page in reader.pages]
    return "\n\n".join(page for page in pages if page)


def _extract_docx(data: bytes) -> str:
    doc = Document(BytesIO(data))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    table_rows: list[str] = []
    for table in doc.tables:
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if values:
                table_rows.append(" | ".join(values))
    return "\n".join([*paragraphs, *table_rows])
